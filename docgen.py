"""
IGNOU Award/Grade List document generator.
One page per course — each course gets its own full IGNOU form page.
"""
import io
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────

def _tc_is_continuation(tc):
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return False
    vm = tcPr.find(qn('w:vMerge'))
    return vm is not None and vm.get(qn('w:val')) is None


def _get_or_create_tcPr(tc):
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    return tcPr


def apply_borders(tc, sz='4', color='000000'):
    if _tc_is_continuation(tc):
        return
    tcPr = _get_or_create_tcPr(tc)
    for ex in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(ex)
    tcB = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        t = OxmlElement(f'w:{edge}')
        t.set(qn('w:val'),   'single')
        t.set(qn('w:sz'),    sz)
        t.set(qn('w:space'), '0')
        t.set(qn('w:color'), color)
        tcB.append(t)
    shd = tcPr.find(qn('w:shd'))
    va  = tcPr.find(qn('w:vAlign'))
    if shd is not None:
        shd.addprevious(tcB)
    elif va is not None:
        va.addprevious(tcB)
    else:
        tcPr.append(tcB)


def apply_bg(tc, fill):
    if _tc_is_continuation(tc):
        return
    tcPr = _get_or_create_tcPr(tc)
    for ex in tcPr.findall(qn('w:shd')):
        tcPr.remove(ex)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    va = tcPr.find(qn('w:vAlign'))
    if va is not None:
        va.addprevious(shd)
    else:
        tcPr.append(shd)


def set_valign(tc, val='center'):
    if _tc_is_continuation(tc):
        return
    tcPr = _get_or_create_tcPr(tc)
    for ex in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(ex)
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), val)
    tcPr.append(va)


def set_cell_text(tc, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    if _tc_is_continuation(tc):
        return
    set_valign(tc)
    p_el = tc.find(qn('w:p'))
    if p_el is None:
        p_el = OxmlElement('w:p')
        tc.append(p_el)
    for r in p_el.findall(qn('w:r')):
        p_el.remove(r)
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_el.insert(0, pPr)
    for tag in (qn('w:spacing'), qn('w:jc')):
        for ex in pPr.findall(tag):
            pPr.remove(ex)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '20')
    sp.set(qn('w:after'),  '20')
    pPr.append(sp)
    jc_map = {
        WD_ALIGN_PARAGRAPH.CENTER: 'center',
        WD_ALIGN_PARAGRAPH.LEFT:   'left',
        WD_ALIGN_PARAGRAPH.RIGHT:  'right',
    }
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), jc_map.get(align, 'center'))
    pPr.append(jc)
    r_el = OxmlElement('w:r')
    rPr  = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    fnt = OxmlElement('w:rFonts')
    fnt.set(qn('w:ascii'), 'Times New Roman')
    fnt.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(fnt)
    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz_el)
    r_el.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = text
    if text != text.strip():
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_el.append(t_el)
    p_el.append(r_el)


def set_col_width(tc, dxa):
    tcPr = _get_or_create_tcPr(tc)
    for ex in tcPr.findall(qn('w:tcW')):
        tcPr.remove(ex)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)


# ─────────────────────────────────────────────────────────────
# Single-page form builder (adds content into an existing doc)
# ─────────────────────────────────────────────────────────────

ROWS_PER_PAGE = 25   # candidates per physical page

def _build_form_page(doc, candidates, course_code, session_label, page_num=1, total_pages=1):
    """
    Append one complete IGNOU Award/Grade List form into `doc`.
    Every physical page gets the full header + table + signature block.
    session_label : e.g. 'Jun 2024' or 'Dec 2024'
    page_num      : 1-based page number within this course (shown when total_pages > 1)
    total_pages   : total pages for this course
    """
    CW = 9638  # content width DXA for A4 with 2 cm margins

    def para(text='', bold=False, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
             sb=0, sa=2, underline=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        if text:
            r = p.add_run(text)
            r.bold = bold; r.underline = underline
            r.font.size = Pt(size); r.font.name = 'Times New Roman'
        return p

    def info_line(lbl1, val1, lbl2, val2, sa=2):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(sa)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(3.35))
        for txt, ul in [(lbl1, False), (val1, True), ('\t', False), (lbl2, False), (val2, True)]:
            r = p.add_run(txt)
            r.underline = ul
            r.font.size = Pt(9)
            r.font.name = 'Times New Roman'

    # ── HEADER ──
    para('INDIRA GANDHI NATIONAL OPEN UNIVERSITY', bold=True, size=13, sa=1)
    para('Study Centre & Code : Naval Headquarters, DNE, RK Puram , New Delhi (Code -7101)',
         size=9, sa=2)
    para('AWARD/GRADE LIST FOR ASSIGNMENTS', bold=True, size=11, underline=True, sa=1)
    session_line = f'For TEE {session_label} Session'
    if total_pages > 1:
        session_line += f'   (Page {page_num} of {total_pages})'
    para(session_line, size=9, sa=4)

    info_line('Programme: ', '', 'Course Code: ', course_code)
    info_line('Study Centre: ', '7101, NAVY', 'Assignment No: ', '_' * 10)
    info_line('Place: ', 'WEST BLOK - V, PORTA CABIN, SECTOR-1, RK PURAM',
              '  For MA Maximum Marks: ', '100', sa=4)
    para('Please arrange Enrolment Nos. in ascending order only and write', size=9, sa=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    para('Complete and correct enrolment number in nine digits.', size=9, sa=4, align=WD_ALIGN_PARAGRAPH.LEFT)

    # ── TABLE ──
    raw_w = [520, 1480, 2480, 1440, 860, 860, 860, 1138]
    scale = CW / sum(raw_w)
    COL_W = [int(w * scale) for w in raw_w]
    COL_W[-1] = CW - sum(COL_W[:-1])

    data_rows  = ROWS_PER_PAGE   # always exactly 25 rows — candidates fill from top, rest blank
    total_rows = 2 + data_rows

    tbl = doc.add_table(rows=total_rows, cols=8)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            set_col_width(cell._tc, COL_W[ci])

    # Merges first
    tbl.rows[0].cells[4].merge(tbl.rows[0].cells[6])
    for col in [0, 1, 2, 3, 7]:
        tbl.rows[0].cells[col].merge(tbl.rows[1].cells[col])

    rows_xml = tbl._tbl.findall(qn('w:tr'))
    r0_tcs   = rows_xml[0].findall(qn('w:tc'))
    r1_tcs   = rows_xml[1].findall(qn('w:tc'))

    HBG   = 'C0CCE0'
    HSIZE = 8

    for tc, text, align in [
        (r0_tcs[0], 'Sl.\nNo',           WD_ALIGN_PARAGRAPH.CENTER),
        (r0_tcs[1], 'Enrolment No.',      WD_ALIGN_PARAGRAPH.CENTER),
        (r0_tcs[2], 'Name of Candidate',  WD_ALIGN_PARAGRAPH.CENTER),
        (r0_tcs[3], 'Programme',          WD_ALIGN_PARAGRAPH.CENTER),
        (r0_tcs[4], 'Grade/Award',        WD_ALIGN_PARAGRAPH.CENTER),
        (r0_tcs[5], 'Remarks\nif any',    WD_ALIGN_PARAGRAPH.CENTER),
    ]:
        set_cell_text(tc, text, bold=True, size=HSIZE, align=align)
        apply_borders(tc, sz='6')
        apply_bg(tc, HBG)

    for tc, text in [(r1_tcs[4], 'TMA-I'), (r1_tcs[5], 'TMA-II'), (r1_tcs[6], 'TMA-III')]:
        set_cell_text(tc, text, bold=True, size=HSIZE)
        apply_borders(tc, sz='6')
        apply_bg(tc, HBG)

    aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]

    for i in range(data_rows):
        row_xml = rows_xml[2 + i]
        trPr = row_xml.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            row_xml.insert(0, trPr)
        for ex in trPr.findall(qn('w:trHeight')):
            trPr.remove(ex)
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), '240')
        trPr.append(trH)

        row_tcs = row_xml.findall(qn('w:tc'))
        if i < len(candidates):
            c    = candidates[i]
            vals = [str(i + 1), c['enrollment'], c['name'], c['programme'], '', '', '', '']
        else:
            vals = [str(i + 1), '', '', '', '', '', '', '']

        for j, (tc, val) in enumerate(zip(row_tcs, vals)):
            set_cell_text(tc, val, size=9, align=aligns[j])
            apply_borders(tc, sz='4')

    # ── SIGNATURE BLOCK ──
    p_gap = doc.add_paragraph()
    p_gap.paragraph_format.space_before = Pt(10)
    p_gap.paragraph_format.space_after  = Pt(0)

    sig_lines = [
        ('Signature of Co-ordinator________________', 'Signature of Evaluation________________'),
        ('Date____________________',                  'Date____________________'),
        ('Office Stamp',                              'Name & Address_________________________________'),
    ]
    for left, right in sig_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(3.35))
        for txt in [left, '\t', right]:
            r = p.add_run(txt)
            r.font.size = Pt(9)
            r.font.name = 'Times New Roman'


def _add_page_break_before(doc):
    """
    Insert a hard page break so the NEXT course's full header
    starts at the top of a fresh page.
    Must be called BEFORE _build_form_page, not after.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    br  = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def _make_doc_skeleton():
    """Create a fresh Document with A4 margins."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    return doc


def _sort_candidates(candidates):
    def enroll_key(c):
        n = ''.join(filter(str.isdigit, str(c.get('enrollment', ''))))
        return int(n) if n else 0
    return sorted(candidates, key=enroll_key)


# ─────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────

def _append_course_pages(doc, course_code, candidates, session_label, is_first):
    """
    Add all physical pages for one course into `doc`.
    Candidates are split into chunks of ROWS_PER_PAGE.
    Every chunk gets a full header + table + signature block.
    A page-break is prepended before every page except the very first
    element in the document.
    """
    sorted_cands = _sort_candidates(candidates)
    # Always produce at least one page even for empty candidate lists
    chunks = [sorted_cands[i:i + ROWS_PER_PAGE]
              for i in range(0, max(1, len(sorted_cands)), ROWS_PER_PAGE)]
    if not chunks:
        chunks = [[]]
    total_pages = len(chunks)

    for page_num, chunk in enumerate(chunks, start=1):
        # Insert a page-break before every page except the very first in the doc
        if not (is_first and page_num == 1):
            _add_page_break_before(doc)
        _build_form_page(doc, chunk, course_code, session_label, page_num, total_pages)


def generate_award_list(course_candidates_map, session_label='Dec 2024'):
    """
    Generate one .docx containing all courses.
    Every course paginates at ROWS_PER_PAGE (25) rows.
    Every physical page — including continuation pages for large courses —
    carries the full IGNOU header + table + signature block.

    course_candidates_map : dict  {course_code: [candidate dicts]}
    session_label         : 'Jun 2024'  or  'Dec 2024'
    Returns               : BytesIO of the .docx
    """
    doc     = _make_doc_skeleton()
    courses = list(course_candidates_map.items())

    for idx, (course_code, candidates) in enumerate(courses):
        _append_course_pages(doc, course_code, candidates, session_label, is_first=(idx == 0))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_award_list_batched(course_candidates_map, session_label='Dec 2024',
                                batch_size=10):
    """
    Generator — yields (batch_number, total_batches, BytesIO) one batch at a time.

    Each batch contains up to `batch_size` courses.  Within each course,
    candidates are paginated at ROWS_PER_PAGE rows and every physical page
    gets the full IGNOU header + table + signature — no bare-entry pages.

    Only ONE Document object exists in RAM at a time.  The caller should
    consume/send each yielded BytesIO before advancing the generator so that
    the previous doc can be garbage-collected.
    """
    courses   = list(course_candidates_map.items())
    n_batches = max(1, (len(courses) + batch_size - 1) // batch_size)

    for batch_idx in range(n_batches):
        batch     = courses[batch_idx * batch_size:(batch_idx + 1) * batch_size]
        doc       = _make_doc_skeleton()

        for i, (course_code, candidates) in enumerate(batch):
            _append_course_pages(doc, course_code, candidates, session_label, is_first=(i == 0))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        yield batch_idx + 1, n_batches, buf
        # Explicitly release both the Document and the BytesIO
        del doc
        buf.close()
